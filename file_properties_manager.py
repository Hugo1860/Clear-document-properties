#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Windows文件属性管理器
支持查看和清除图片、PDF、Word文档的属性信息
"""

import os
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image
from PIL.ExifTags import TAGS
import fitz  # PyMuPDF
import docx
import win32com.client
import pythoncom
import threading
import traceback

class FilePropertiesManager:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Windows文件属性管理器")
        self.root.geometry("800x600")
        
        # 设置窗口图标和样式
        self.root.configure(bg='#f0f0f0')
        
        # 创建主框架
        self.create_widgets()
        
        # 当前选中的文件路径
        self.current_file = None
        
    def create_widgets(self):
        # 顶部框架
        top_frame = ttk.Frame(self.root, padding="10")
        top_frame.pack(fill=tk.X)
        
        # 文件选择区域
        file_label = ttk.Label(top_frame, text="选择文件:")
        file_label.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        
        self.file_path_var = tk.StringVar()
        file_entry = ttk.Entry(top_frame, textvariable=self.file_path_var, width=60)
        file_entry.grid(row=0, column=1, padx=5, pady=5)
        
        browse_btn = ttk.Button(top_frame, text="浏览", command=self.browse_file)
        browse_btn.grid(row=0, column=2, padx=5, pady=5)
        
        # 按钮区域
        btn_frame = ttk.Frame(top_frame)
        btn_frame.grid(row=1, column=0, columnspan=3, pady=10)
        
        view_btn = ttk.Button(btn_frame, text="查看属性", command=self.view_properties)
        view_btn.pack(side=tk.LEFT, padx=5)
        
        clear_btn = ttk.Button(btn_frame, text="清除属性", command=self.clear_properties)
        clear_btn.pack(side=tk.LEFT, padx=5)
        
        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='indeterminate')
        self.progress.pack(fill=tk.X, padx=10, pady=5)
        
        # 属性显示区域
        self.create_properties_display()
        
    def create_properties_display(self):
        # 创建Notebook用于分类显示属性
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 基本属性页面
        self.basic_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.basic_frame, text="基本属性")
        
        # EXIF属性页面
        self.exif_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.exif_frame, text="EXIF信息")
        
        # PDF属性页面
        self.pdf_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.pdf_frame, text="PDF属性")
        
        # Word属性页面
        self.word_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.word_frame, text="Word属性")
        
        # 为每个页面创建文本框
        for frame in [self.basic_frame, self.exif_frame, self.pdf_frame, self.word_frame]:
            text_widget = tk.Text(frame, wrap=tk.WORD, width=80, height=20)
            scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=text_widget.yview)
            text_widget.configure(yscrollcommand=scrollbar.set)
            
            text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            
            # 保存引用
            if frame == self.basic_frame:
                self.basic_text = text_widget
            elif frame == self.exif_frame:
                self.exif_text = text_widget
            elif frame == self.pdf_frame:
                self.pdf_text = text_widget
            elif frame == self.word_frame:
                self.word_text = text_widget
    
    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="选择文件",
            filetypes=[
                ("支持的文件", "*.jpg *.jpeg *.png *.gif *.bmp *.pdf *.docx *.doc"),
                ("图片文件", "*.jpg *.jpeg *.png *.gif *.bmp"),
                ("PDF文件", "*.pdf"),
                ("Word文件", "*.docx *.doc"),
                ("所有文件", "*.*")
            ]
        )
        if filename:
            self.file_path_var.set(filename)
            self.current_file = filename
    
    def view_properties(self):
        if not self.current_file:
            messagebox.showwarning("警告", "请先选择文件！")
            return
        
        if not os.path.exists(self.current_file):
            messagebox.showerror("错误", "文件不存在！")
            return
        
        # 在新线程中处理，避免界面卡顿
        threading.Thread(target=self._load_properties, daemon=True).start()
    
    def _load_properties(self):
        try:
            self.root.after(0, lambda: self.progress.start())
            
            # 清空所有显示
            self.root.after(0, lambda: self.clear_all_displays())
            
            # 获取基本属性
            basic_info = self.get_basic_properties()
            self.root.after(0, lambda: self.basic_text.insert(1.0, basic_info))
            
            # 根据文件类型获取特定属性
            file_ext = os.path.splitext(self.current_file)[1].lower()
            
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                exif_info = self.get_image_exif()
                self.root.after(0, lambda: self.exif_text.insert(1.0, exif_info))
                self.root.after(0, lambda: self.notebook.select(self.exif_frame))
            elif file_ext == '.pdf':
                pdf_info = self.get_pdf_properties()
                self.root.after(0, lambda: self.pdf_text.insert(1.0, pdf_info))
                self.root.after(0, lambda: self.notebook.select(self.pdf_frame))
            elif file_ext in ['.docx', '.doc']:
                word_info = self.get_word_properties()
                self.root.after(0, lambda: self.word_text.insert(1.0, word_info))
                self.root.after(0, lambda: self.notebook.select(self.word_frame))
            
        except Exception as e:
            error_msg = f"加载属性时出错: {str(e)}\n{traceback.format_exc()}"
            self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
        finally:
            self.root.after(0, lambda: self.progress.stop())
    
    def clear_all_displays(self):
        for text_widget in [self.basic_text, self.exif_text, self.pdf_text, self.word_text]:
            text_widget.delete(1.0, tk.END)
    
    def get_basic_properties(self):
        try:
            stat = os.stat(self.current_file)
            info = []
            info.append(f"文件名: {os.path.basename(self.current_file)}")
            info.append(f"完整路径: {self.current_file}")
            info.append(f"文件大小: {self.format_file_size(stat.st_size)}")
            info.append(f"创建时间: {self.format_timestamp(stat.st_ctime)}")
            info.append(f"修改时间: {self.format_timestamp(stat.st_mtime)}")
            info.append(f"访问时间: {self.format_timestamp(stat.st_atime)}")
            
            # Windows特定属性
            if os.name == 'nt':
                try:
                    import win32api
                    import win32con
                    attrs = win32api.GetFileAttributes(self.current_file)
                    info.append(f"文件属性: {self.get_file_attributes(attrs)}")
                except:
                    pass
            
            return "\n".join(info)
        except Exception as e:
            return f"获取基本属性失败: {str(e)}"
    
    def get_image_exif(self):
        try:
            image = Image.open(self.current_file)
            exifdata = image.getexif()
            
            if not exifdata:
                return "该图片没有EXIF信息"
            
            info = []
            for tag_id in exifdata:
                tag = TAGS.get(tag_id, tag_id)
                data = exifdata.get(tag_id)
                
                # 处理二进制数据
                if isinstance(data, bytes):
                    try:
                        data = data.decode('utf-8', errors='ignore')
                    except:
                        data = f"<二进制数据: {len(data)}字节>"
                
                info.append(f"{tag}: {data}")
            
            return "\n".join(info) if info else "没有EXIF信息"
        except Exception as e:
            return f"获取EXIF信息失败: {str(e)}"
    
    def get_pdf_properties(self):
        try:
            doc = fitz.open(self.current_file)
            info = []
            
            # 获取PDF版本信息
            metadata = doc.metadata
            pdf_version = metadata.get('format', '未知') if metadata else '未知'
            info.append(f"PDF版本: {pdf_version}")
            info.append(f"页面数量: {len(doc)}")
            
            # 元数据
            if metadata:
                info.append("\n元数据:")
                for key, value in metadata.items():
                    if value:
                        info.append(f"  {key}: {value}")
            
            # 页面信息
            info.append("\n页面信息:")
            for i, page in enumerate(doc):
                rect = page.rect
                info.append(f"  页面 {i+1}: {rect.width}x{rect.height} 点")
            
            doc.close()
            return "\n".join(info)
        except Exception as e:
            return f"获取PDF属性失败: {str(e)}"
    
    def get_word_properties(self):
        try:
            if self.current_file.lower().endswith('.docx'):
                doc = docx.Document(self.current_file)
                
                info = []
                props = doc.core_properties
                
                info.append("文档属性:")
                info.append(f"  标题: {props.title or '无'}")
                info.append(f"  主题: {props.subject or '无'}")
                info.append(f"  作者: {props.author or '无'}")
                info.append(f"  类别: {props.category or '无'}")
                info.append(f"  关键词: {props.keywords or '无'}")
                info.append(f"  备注: {props.comments or '无'}")
                info.append(f"  最后修改者: {props.last_modified_by or '无'}")
                info.append(f"  修订号: {props.revision or '无'}")
                
                info.append(f"\n  创建时间: {props.created}")
                info.append(f"  最后修改时间: {props.modified}")
                info.append(f"  最后打印时间: {props.last_printed}")
                
                # 统计信息
                paragraphs = len(doc.paragraphs)
                tables = len(doc.tables)
                
                info.append(f"\n统计信息:")
                info.append(f"  段落数: {paragraphs}")
                info.append(f"  表格数: {tables}")
                
                return "\n".join(info)
            else:  # .doc文件
                return "DOC格式需要安装Microsoft Word才能查看详细属性"
        except Exception as e:
            return f"获取Word属性失败: {str(e)}"
    
    def clear_properties(self):
        if not self.current_file:
            messagebox.showwarning("警告", "请先选择文件！")
            return
        
        if not os.path.exists(self.current_file):
            messagebox.showerror("错误", "文件不存在！")
            return
        
        if messagebox.askyesno("确认", "确定要清除此文件的所有属性信息吗？此操作不可撤销！"):
            threading.Thread(target=self._clear_properties, daemon=True).start()
    
    def _clear_properties(self):
        try:
            self.root.after(0, lambda: self.progress.start())
            
            file_ext = os.path.splitext(self.current_file)[1].lower()
            
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                self.clear_image_properties()
            elif file_ext == '.pdf':
                self.clear_pdf_properties()
            elif file_ext in ['.docx', '.doc']:
                self.clear_word_properties()
            else:
                self.root.after(0, lambda: messagebox.showwarning("警告", "不支持的文件类型"))
                return
            
            self.root.after(0, lambda: messagebox.showinfo("成功", "属性清除完成！"))
            self.root.after(0, lambda: self.view_properties())  # 重新加载属性
            
        except Exception as e:
            error_msg = f"清除属性时出错: {str(e)}\n{traceback.format_exc()}"
            self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
        finally:
            self.root.after(0, lambda: self.progress.stop())
    
    def clear_image_properties(self):
        try:
            # 创建新图片，不包含EXIF信息
            image = Image.open(self.current_file)
            
            # 保存为新文件，不包含EXIF
            temp_path = self.current_file + '.tmp'
            image.save(temp_path, quality=95)
            
            # 替换原文件
            os.replace(temp_path, self.current_file)
            
        except Exception as e:
            raise Exception(f"清除图片属性失败: {str(e)}")
    
    def clear_pdf_properties(self):
        try:
            doc = fitz.open(self.current_file)
            
            # 清除元数据
            doc.set_metadata({})
            
            # 保存到新文件
            temp_path = self.current_file + '.tmp'
            doc.save(temp_path)
            doc.close()
            
            # 替换原文件
            os.replace(temp_path, self.current_file)
            
        except Exception as e:
            raise Exception(f"清除PDF属性失败: {str(e)}")
    
    def clear_word_properties(self):
        try:
            if self.current_file.lower().endswith('.docx'):
                doc = docx.Document(self.current_file)
                
                # 清除核心属性
                props = doc.core_properties
                props.title = None
                props.subject = None
                props.author = None
                props.category = None
                props.keywords = None
                props.comments = None
                props.last_modified_by = None
                
                # 保存到新文件
                temp_path = self.current_file + '.tmp'
                doc.save(temp_path)
                
                # 替换原文件
                os.replace(temp_path, self.current_file)
            else:
                # 对于DOC文件，使用COM接口清除属性
                self.clear_doc_properties_com()
                
        except Exception as e:
            raise Exception(f"清除Word属性失败: {str(e)}")
    
    def clear_doc_properties_com(self):
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(self.current_file)
            
            # 清除内置文档属性
            props = doc.BuiltInDocumentProperties
            for prop in props:
                try:
                    if prop.Name in ["Title", "Subject", "Author", "Keywords", "Comments"]:
                        prop.Value = ""
                except:
                    pass
            
            # 保存并关闭
            doc.Save()
            doc.Close()
            word.Quit()
            
            pythoncom.CoUninitialize()
            
        except Exception as e:
            if 'pythoncom' in locals():
                pythoncom.CoUninitialize()
            raise Exception(f"使用Word COM清除属性失败: {str(e)}")
    
    def format_file_size(self, size):
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.2f} {unit}"
            size /= 1024
        return f"{size:.2f} TB"
    
    def format_timestamp(self, timestamp):
        import datetime
        return datetime.datetime.fromtimestamp(timestamp).strftime('%Y-%m-%d %H:%M:%S')
    
    def get_file_attributes(self, attrs):
        attributes = []
        if attrs & 0x1: attributes.append("只读")
        if attrs & 0x2: attributes.append("隐藏")
        if attrs & 0x4: attributes.append("系统")
        if attrs & 0x10: attributes.append("目录")
        if attrs & 0x20: attributes.append("存档")
        return ", ".join(attributes) if attributes else "无特殊属性"
    
    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = FilePropertiesManager()
    app.run()