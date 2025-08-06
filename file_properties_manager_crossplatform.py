#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
跨平台文件属性管理器
支持Windows、macOS、Linux系统
支持查看和清除图片、PDF、Word文档的属性信息
"""

import os
import sys
import platform
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image
from PIL.ExifTags import TAGS
import fitz  # PyMuPDF
import docx
import threading
import traceback
import subprocess

# 平台检测
SYSTEM = platform.system()
IS_WINDOWS = SYSTEM == "Windows"
IS_MACOS = SYSTEM == "Darwin"
IS_LINUX = SYSTEM == "Linux"

class FilePropertiesManager:
    def __init__(self):
        self.root = tk.Tk()
        
        # 设置窗口标题根据平台
        if IS_WINDOWS:
            self.root.title("Windows文件属性管理器 - 批量版")
        elif IS_MACOS:
            self.root.title("macOS文件属性管理器 - 批量版")
        else:
            self.root.title("跨平台文件属性管理器 - 批量版")
            
        self.root.geometry("1000x700")
        self.root.configure(bg='#f0f0f0')
        
        # 文件列表
        self.file_list = []
        self.selected_files = []
        
        # 创建主框架
        self.create_widgets()
        
        # 当前选中的文件路径（单文件模式）
        self.current_file = None
        
    def create_widgets(self):
        # 顶部框架
        top_frame = ttk.Frame(self.root, padding="10")
        top_frame.pack(fill=tk.X)

        # 文件选择区域
        file_label = ttk.Label(top_frame, text="选择文件:")
        file_label.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)

        self.file_path_var = tk.StringVar()
        file_entry = ttk.Entry(top_frame, textvariable=self.file_path_var, width=60)
        file_entry.grid(row=0, column=1, padx=5, pady=5)

        browse_btn = ttk.Button(top_frame, text="浏览", command=self.browse_file)
        browse_btn.grid(row=0, column=2, padx=5, pady=5)

        # 按钮区域
        btn_frame = ttk.Frame(top_frame)
        btn_frame.grid(row=1, column=0, columnspan=3, pady=10)

        view_btn = ttk.Button(btn_frame, text="查看属性", command=self.view_properties)
        view_btn.pack(side=tk.LEFT, padx=5)

        clear_btn = ttk.Button(btn_frame, text="清除属性", command=self.clear_properties)
        clear_btn.pack(side=tk.LEFT, padx=5)

        # 批量处理按钮区域
        batch_btn_frame = ttk.Frame(top_frame)
        batch_btn_frame.grid(row=2, column=0, columnspan=3, pady=5)

        ttk.Button(batch_btn_frame, text="批量添加文件", command=self.batch_add_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(batch_btn_frame, text="批量查看属性", command=self.batch_view_properties).pack(side=tk.LEFT, padx=5)
        ttk.Button(batch_btn_frame, text="一键批量清除", command=self.batch_clear_properties).pack(side=tk.LEFT, padx=5)
        ttk.Button(batch_btn_frame, text="移除已处理文件", command=self.remove_processed_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(batch_btn_frame, text="全选", command=self.select_all_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(batch_btn_frame, text="取消全选", command=self.deselect_all_files).pack(side=tk.LEFT, padx=5)

        # 文件列表区域
        self.create_file_list_area()

        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='determinate', length=300)
        self.progress.pack(fill=tk.X, padx=10, pady=5)

        # 属性显示区域
        self.create_properties_display()

        # 显示系统信息
        self.show_system_info()
        
    def create_file_list_area(self):
        """创建文件列表区域"""
        # 文件列表标题
        title_frame = ttk.Frame(self.root)
        title_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(title_frame, text="文件列表", font=('Arial', 12, 'bold')).pack(side=tk.LEFT)
        ttk.Label(title_frame, text="(支持批量选择)").pack(side=tk.LEFT, padx=10)
        
        # 文件列表框架
        list_frame = ttk.Frame(self.root)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 创建Treeview显示文件列表
        columns = ('文件名', '类型', '大小', '路径')
        self.file_tree = ttk.Treeview(list_frame, columns=columns, show='tree headings', height=15)
        
        # 定义列
        self.file_tree.heading('#0', text='☐')
        self.file_tree.heading('文件名', text='文件名')
        self.file_tree.heading('类型', text='类型')
        self.file_tree.heading('大小', text='大小')
        self.file_tree.heading('路径', text='完整路径')
        
        # 设置列宽
        self.file_tree.column('#0', width=60, minwidth=60)
        self.file_tree.column('文件名', width=200, minwidth=120)
        self.file_tree.column('类型', width=80, minwidth=60)
        self.file_tree.column('大小', width=100, minwidth=80)
        self.file_tree.column('路径', width=300, minwidth=150)
        
        # 设置选中行的样式
        self.file_tree.tag_configure('selected', background='#e6f3ff')
        
        # 滚动条
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
        self.file_tree.configure(yscrollcommand=scrollbar.set)
        
        self.file_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 绑定单击事件实现复选框效果
        self.file_tree.bind('<Button-1>', self.on_file_click)
        
    def create_properties_display(self):
        """创建属性显示区域"""
        # 创建Notebook用于分类显示属性
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 系统信息页面
        self.system_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.system_frame, text="系统信息")

        # 批量处理结果页面
        self.batch_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.batch_frame, text="批量处理结果")

        # 基本属性页面
        self.basic_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.basic_frame, text="基本属性")

        # EXIF属性页面
        self.exif_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.exif_frame, text="EXIF信息")

        # PDF属性页面
        self.pdf_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.pdf_frame, text="PDF属性")

        # Word属性页面
        self.word_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.word_frame, text="Word属性")

        # 为每个页面创建文本框
        frames = [self.system_frame, self.batch_frame, self.basic_frame, self.exif_frame, self.pdf_frame, self.word_frame]
        text_widgets = []

        for frame in frames:
            text_widget = tk.Text(frame, wrap=tk.WORD, width=80, height=25)
            scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=text_widget.yview)
            text_widget.configure(yscrollcommand=scrollbar.set)

            text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            text_widgets.append(text_widget)

        # 保存引用
        self.system_text, self.batch_text, self.basic_text, self.exif_text, self.pdf_text, self.word_text = text_widgets
        
    def show_system_info(self):
        info = []
        info.append(f"操作系统: {platform.system()} {platform.release()}")
        info.append(f"Python版本: {platform.python_version()}")
        info.append(f"处理器: {platform.processor()}")
        info.append(f"架构: {platform.machine()}")
        info.append(f"工作目录: {os.getcwd()}")
        
        self.system_text.insert(1.0, "\n".join(info))
        
    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="选择文件",
            filetypes=[
                ("支持的文件", "*.jpg *.jpeg *.png *.gif *.bmp *.pdf *.docx *.doc"),
                ("图片文件", "*.jpg *.jpeg *.png *.gif *.bmp"),
                ("PDF文件", "*.pdf"),
                ("Word文件", "*.docx *.doc"),
                ("所有文件", "*.*")
            ]
        )
        if filename:
            self.file_path_var.set(filename)
            self.current_file = filename
            
    def batch_add_files(self):
        """批量添加文件"""
        filenames = filedialog.askopenfilenames(
            title="批量选择文件",
            filetypes=[
                ("支持的文件", "*.jpg *.jpeg *.png *.gif *.bmp *.pdf *.docx *.doc"),
                ("图片文件", "*.jpg *.jpeg *.png *.gif *.bmp"),
                ("PDF文件", "*.pdf"),
                ("Word文件", "*.docx *.doc"),
                ("所有文件", "*.*")
            ]
        )
        
        for filename in filenames:
            if filename not in self.file_list:
                self.file_list.append(filename)
                self.add_file_to_tree(filename)
                
        self.update_batch_status()
        
    def add_file_to_tree(self, filepath):
        """添加文件到Treeview"""
        try:
            filename = os.path.basename(filepath)
            file_ext = os.path.splitext(filepath)[1].lower()
            file_size = os.path.getsize(filepath)
            
            # 确定文件类型
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                file_type = "图片"
            elif file_ext == '.pdf':
                file_type = "PDF"
            elif file_ext in ['.docx', '.doc']:
                file_type = "Word"
            else:
                file_type = "其他"
                
            # 插入到Treeview
            self.file_tree.insert('', 'end', text='☐', values=(
                filename,
                file_type,
                self.format_file_size(file_size),
                filepath
            ))
            
        except Exception as e:
            messagebox.showerror("错误", f"添加文件失败: {str(e)}")
            
    def select_all_files(self):
        """全选文件"""
        for item in self.file_tree.get_children():
            self.file_tree.item(item, tags=['selected'])
            self.file_tree.item(item, text='✓')
            
    def deselect_all_files(self):
        """取消全选"""
        for item in self.file_tree.get_children():
            self.file_tree.item(item, tags=[])
            self.file_tree.item(item, text='')
            
    def on_file_click(self, event):
        """单击文件切换选择状态"""
        region = self.file_tree.identify_region(event.x, event.y)
        if region == "tree":  # 点击了选择列
            item = self.file_tree.identify_row(event.y)
            if item:
                tags = list(self.file_tree.item(item, 'tags'))
                current_text = self.file_tree.item(item, 'text')
                
                if 'selected' in tags:
                    tags.remove('selected')
                    self.file_tree.item(item, tags=tags)
                    self.file_tree.item(item, text='☐')
                else:
                    tags.append('selected')
                    self.file_tree.item(item, tags=tags)
                    self.file_tree.item(item, text='☑')
                return "break"  # 阻止默认行为
                
    def get_selected_files(self):
        """获取选中的文件列表"""
        selected_files = []
        for item in self.file_tree.get_children():
            tags = self.file_tree.item(item, 'tags')
            if tags and 'selected' in tags:
                values = self.file_tree.item(item, 'values')
                if values and len(values) > 3:
                    selected_files.append(values[3])  # 路径在最后一列
        return selected_files
            
    def view_properties(self):
        if not self.current_file:
            messagebox.showwarning("警告", "请先选择文件！")
            return
        
        if not os.path.exists(self.current_file):
            messagebox.showerror("错误", "文件不存在！")
            return
        
        threading.Thread(target=self._load_properties, daemon=True).start()
        
    def _load_properties(self):
        try:
            self.root.after(0, lambda: self.progress.start())
            self.root.after(0, lambda: self.clear_all_displays())
            
            basic_info = self.get_basic_properties()
            self.root.after(0, lambda: self.basic_text.insert(1.0, basic_info))
            
            file_ext = os.path.splitext(self.current_file)[1].lower()
            
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                exif_info = self.get_image_exif()
                self.root.after(0, lambda: self.exif_text.insert(1.0, exif_info))
                self.root.after(0, lambda: self.notebook.select(self.exif_frame))
            elif file_ext == '.pdf':
                pdf_info = self.get_pdf_properties()
                self.root.after(0, lambda: self.pdf_text.insert(1.0, pdf_info))
                self.root.after(0, lambda: self.notebook.select(self.pdf_frame))
            elif file_ext in ['.docx', '.doc']:
                word_info = self.get_word_properties()
                self.root.after(0, lambda: self.word_text.insert(1.0, word_info))
                self.root.after(0, lambda: self.notebook.select(self.word_frame))
                
        except Exception as e:
            error_msg = f"加载属性时出错: {str(e)}\n{traceback.format_exc()}"
            self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
        finally:
            self.root.after(0, lambda: self.progress.stop())
            
    def clear_all_displays(self):
        for text_widget in [self.system_text, self.basic_text, self.exif_text, self.pdf_text, self.word_text]:
            text_widget.delete(1.0, tk.END)
            
    def get_basic_properties(self):
        try:
            stat = os.stat(self.current_file)
            info = []
            info.append(f"文件名: {os.path.basename(self.current_file)}")
            info.append(f"完整路径: {self.current_file}")
            info.append(f"文件大小: {self.format_file_size(stat.st_size)}")
            info.append(f"创建时间: {self.format_timestamp(stat.st_ctime)}")
            info.append(f"修改时间: {self.format_timestamp(stat.st_mtime)}")
            info.append(f"访问时间: {self.format_timestamp(stat.st_atime)}")
            
            # 平台特定属性
            if IS_WINDOWS:
                try:
                    import win32api
                    import win32con
                    attrs = win32api.GetFileAttributes(self.current_file)
                    info.append(f"文件属性: {self.get_file_attributes(attrs)}")
                except:
                    pass
            elif IS_MACOS:
                # macOS文件属性
                try:
                    import subprocess
                    result = subprocess.run(['ls', '-la', self.current_file], 
                                          capture_output=True, text=True)
                    if result.returncode == 0:
                        parts = result.stdout.split()
                        if len(parts) >= 10:
                            info.append(f"权限: {parts[0]}")
                            info.append(f"所有者: {parts[2]}")
                            info.append(f"群组: {parts[3]}")
                except:
                    pass
            elif IS_LINUX:
                # Linux文件属性
                try:
                    result = subprocess.run(['ls', '-la', self.current_file], 
                                          capture_output=True, text=True)
                    if result.returncode == 0:
                        parts = result.stdout.split()
                        if len(parts) >= 9:
                            info.append(f"权限: {parts[0]}")
                            info.append(f"所有者: {parts[2]}")
                            info.append(f"群组: {parts[3]}")
                except:
                    pass
            
            return "\n".join(info)
        except Exception as e:
            return f"获取基本属性失败: {str(e)}"
            
    def get_image_exif(self):
        try:
            image = Image.open(self.current_file)
            exifdata = image.getexif()
            
            if not exifdata:
                return "该图片没有EXIF信息"
            
            info = []
            for tag_id in exifdata:
                tag = TAGS.get(tag_id, tag_id)
                data = exifdata.get(tag_id)
                
                # 处理二进制数据
                if isinstance(data, bytes):
                    try:
                        data = data.decode('utf-8', errors='ignore')
                    except:
                        data = f"<二进制数据: {len(data)}字节>"
                
                info.append(f"{tag}: {data}")
            
            return "\n".join(info) if info else "没有EXIF信息"
        except Exception as e:
            return f"获取EXIF信息失败: {str(e)}"
            
    def get_pdf_properties(self):
        try:
            doc = fitz.open(self.current_file)
            info = []
            
            # 获取PDF版本信息
            metadata = doc.metadata
            pdf_version = metadata.get('format', '未知') if metadata else '未知'
            info.append(f"PDF版本: {pdf_version}")
            info.append(f"页面数量: {len(doc)}")
            
            # 元数据
            if metadata:
                info.append("\n元数据:")
                for key, value in metadata.items():
                    if value:
                        info.append(f"  {key}: {value}")
            
            # 页面信息
            info.append("\n页面信息:")
            for i, page in enumerate(doc):
                rect = page.rect
                info.append(f"  页面 {i+1}: {rect.width}x{rect.height} 点")
            
            doc.close()
            return "\n".join(info)
        except Exception as e:
            return f"获取PDF属性失败: {str(e)}"
            
    def get_word_properties(self):
        try:
            if self.current_file.lower().endswith('.docx'):
                doc = docx.Document(self.current_file)
                
                info = []
                props = doc.core_properties
                
                info.append("文档属性:")
                info.append(f"  标题: {props.title or '无'}")
                info.append(f"  主题: {props.subject or '无'}")
                info.append(f"  作者: {props.author or '无'}")
                info.append(f"  类别: {props.category or '无'}")
                info.append(f"  关键词: {props.keywords or '无'}")
                info.append(f"  备注: {props.comments or '无'}")
                info.append(f"  最后修改者: {props.last_modified_by or '无'}")
                info.append(f"  修订号: {props.revision or '无'}")
                
                info.append(f"\n  创建时间: {props.created}")
                info.append(f"  最后修改时间: {props.modified}")
                info.append(f"  最后打印时间: {props.last_printed}")
                
                # 统计信息
                paragraphs = len(doc.paragraphs)
                tables = len(doc.tables)
                
                info.append(f"\n统计信息:")
                info.append(f"  段落数: {paragraphs}")
                info.append(f"  表格数: {tables}")
                
                return "\n".join(info)
            else:  # .doc文件
                if IS_WINDOWS:
                    return "DOC格式需要安装Microsoft Word才能查看详细属性"
                else:
                    return "DOC格式在macOS/Linux上支持有限，建议使用DOCX格式"
        except Exception as e:
            return f"获取Word属性失败: {str(e)}"
            
    def clear_properties(self):
        if not self.current_file:
            messagebox.showwarning("警告", "请先选择文件！")
            return
        
        if not os.path.exists(self.current_file):
            messagebox.showerror("错误", "文件不存在！")
            return
        
        if messagebox.askyesno("确认", "确定要清除此文件的所有属性信息吗？此操作不可撤销！"):
            threading.Thread(target=self._clear_properties, daemon=True).start()
            
    def batch_view_properties(self):
        """批量查看属性"""
        selected_files = self.get_selected_files()
        if not selected_files:
            messagebox.showwarning("警告", "请先选择要查看的文件！")
            return
            
        threading.Thread(target=self._batch_view_properties_worker, args=(selected_files,), daemon=True).start()
        
    def _batch_view_properties_worker(self, files):
        """批量查看属性工作线程"""
        try:
            self.root.after(0, lambda: self.progress.start())
            self.root.after(0, lambda: self.batch_text.delete(1.0, tk.END))
            
            total_files = len(files)
            results = []
            
            for i, filepath in enumerate(files):
                try:
                    result = self.get_file_summary_info(filepath)
                    results.append(result)
                    
                    # 更新进度
                    progress = (i + 1) / total_files * 100
                    self.root.after(0, lambda p=progress: self.progress.configure(value=p))
                    
                except Exception as e:
                    results.append(f"文件 {filepath} 处理失败: {str(e)}")
            
            # 显示结果
            summary = f"批量查看完成！共处理 {total_files} 个文件\n\n"
            summary += "=" * 60 + "\n\n"
            summary += "\n\n".join(results)
            
            self.root.after(0, lambda: self.batch_text.insert(1.0, summary))
            self.root.after(0, lambda: self.notebook.select(self.batch_frame))
            
        finally:
            self.root.after(0, lambda: self.progress.stop())
            self.root.after(0, lambda: self.progress.configure(value=0))
            
    def batch_clear_properties(self):
        """一键批量清除属性"""
        selected_files = self.get_selected_files()
        if not selected_files:
            messagebox.showwarning("警告", "请先选择要清除的文件！")
            return
            
        if messagebox.askyesno("确认", f"确定要清除 {len(selected_files)} 个文件的所有属性信息吗？\n此操作不可撤销！"):
            threading.Thread(target=self._batch_clear_properties_worker, args=(selected_files,), daemon=True).start()
            
    def _batch_clear_properties_worker(self, files):
        """批量清除属性工作线程"""
        try:
            self.root.after(0, lambda: self.progress.start())
            self.root.after(0, lambda: self.batch_text.delete(1.0, tk.END))
            
            total_files = len(files)
            results = []
            
            for i, filepath in enumerate(files):
                try:
                    result = self.clear_file_properties(filepath)
                    results.append(f"✅ {os.path.basename(filepath)}: {result}")
                    
                    # 更新进度
                    progress = (i + 1) / total_files * 100
                    self.root.after(0, lambda p=progress: self.progress.configure(value=p))
                    
                except Exception as e:
                    results.append(f"❌ {os.path.basename(filepath)}: 失败 - {str(e)}")
            
            # 显示结果
            summary = f"批量清除完成！共处理 {total_files} 个文件\n\n"
            summary += "=" * 60 + "\n\n"
            summary += "\n".join(results)
            
            self.root.after(0, lambda: self.batch_text.insert(1.0, summary))
            self.root.after(0, lambda: self.notebook.select(self.batch_frame))
            
        finally:
            self.root.after(0, lambda: self.progress.stop())
            self.root.after(0, lambda: self.progress.configure(value=0))
            
    def get_file_summary_info(self, filepath):
        """获取文件摘要信息"""
        try:
            filename = os.path.basename(filepath)
            file_ext = os.path.splitext(filepath)[1].lower()
            file_size = os.path.getsize(filepath)
            
            info = [f"文件: {filename}"]
            info.append(f"路径: {filepath}")
            info.append(f"大小: {self.format_file_size(file_size)}")
            
            # 临时保存当前文件，以便其他方法使用
            old_current = self.current_file
            self.current_file = filepath
            
            try:
                if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                    exif_info = self.get_image_exif()
                    if "没有EXIF信息" not in exif_info and "失败" not in exif_info:
                        info.append("状态: 包含EXIF信息")
                    else:
                        info.append("状态: 无EXIF信息")
                        
                elif file_ext == '.pdf':
                    pdf_info = self.get_pdf_properties()
                    if "获取PDF属性失败" not in pdf_info:
                        lines = pdf_info.split('\n')
                        for line in lines[:5]:  # 显示前5行
                            if line.strip():
                                info.append(line)
                    else:
                        info.append("状态: 无法读取PDF属性")
                        
                elif file_ext in ['.docx', '.doc']:
                    word_info = self.get_word_properties()
                    if "获取Word属性失败" not in word_info:
                        lines = word_info.split('\n')
                        for line in lines[:5]:  # 显示前5行
                            if line.strip():
                                info.append(line)
                    else:
                        info.append("状态: 无法读取Word属性")
            finally:
                # 恢复原来的当前文件
                self.current_file = old_current
            
            return "\n".join(info)
            
        except Exception as e:
            return f"文件 {filepath} 处理失败: {str(e)}"
            
    def clear_file_properties(self, filepath):
        """清除单个文件的属性"""
        try:
            file_ext = os.path.splitext(filepath)[1].lower()
            
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                self.clear_image_properties_file(filepath)
                return "图片属性已清除"
            elif file_ext == '.pdf':
                self.clear_pdf_properties_file(filepath)
                return "PDF属性已清除"
            elif file_ext in ['.docx', '.doc']:
                self.clear_word_properties(filepath)
                return "Word属性已清除"
            else:
                return "不支持的文件类型"
                
        except Exception as e:
            raise Exception(f"清除失败: {str(e)}")
            
    def remove_processed_files(self):
        """移除已处理的文件"""
        selected_files = self.get_selected_files()
        if not selected_files:
            messagebox.showwarning("警告", "请先选择要移除的文件！")
            return
            
        if messagebox.askyesno("确认", f"确定要从列表中移除 {len(selected_files)} 个已处理的文件吗？"):
            # 获取选中的文件路径
            selected_paths = set(selected_files)
            
            # 从文件列表中移除
            self.file_list = [f for f in self.file_list if f not in selected_paths]
            
            # 从Treeview中移除
            items_to_remove = []
            for item in self.file_tree.get_children():
                values = self.file_tree.item(item, 'values')
                if values and len(values) > 3 and values[3] in selected_paths:
                    items_to_remove.append(item)
            
            for item in items_to_remove:
                self.file_tree.delete(item)
                
            self.update_batch_status()
            messagebox.showinfo("成功", f"已从列表中移除 {len(items_to_remove)} 个文件")
            
    def update_batch_status(self):
        """更新批量处理状态"""
        total = len(self.file_tree.get_children())
        selected = len(self.get_selected_files())
        status = f"总计: {total} 个文件 | 已选择: {selected} 个文件"
        # 这里可以添加状态栏显示
            
    def _clear_properties(self):
        try:
            self.root.after(0, lambda: self.progress.start())
            
            result = self.clear_file_properties(self.current_file)
            
            self.root.after(0, lambda: messagebox.showinfo("成功", f"属性清除完成！{result}"))
            self.root.after(0, lambda: self.view_properties())  # 重新加载属性
            
        except Exception as e:
            error_msg = f"清除属性时出错: {str(e)}\n{traceback.format_exc()}"
            self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
        finally:
            self.root.after(0, lambda: self.progress.stop())
            
    def clear_image_properties(self, filepath=None):
        try:
            target_file = filepath or self.current_file
            image = Image.open(target_file)
            
            # 保存为新文件，不包含EXIF
            temp_path = target_file + '.tmp'
            image.save(temp_path, quality=95)
            
            # 替换原文件
            os.replace(temp_path, target_file)
            
        except Exception as e:
            raise Exception(f"清除图片属性失败: {str(e)}")
            
    def clear_image_properties_file(self, filepath):
        """清除指定图片文件的属性"""
        old_current = self.current_file
        try:
            self.current_file = filepath
            self.clear_image_properties()
        finally:
            self.current_file = old_current
            
    def clear_pdf_properties(self, filepath=None):
        try:
            target_file = filepath or self.current_file
            doc = fitz.open(target_file)
            
            # 清除元数据
            doc.set_metadata({})
            
            # 保存到新文件
            temp_path = target_file + '.tmp'
            doc.save(temp_path)
            doc.close()
            
            # 替换原文件
            os.replace(temp_path, target_file)
            
        except Exception as e:
            raise Exception(f"清除PDF属性失败: {str(e)}")
            
    def clear_pdf_properties_file(self, filepath):
        """清除指定PDF文件的属性"""
        old_current = self.current_file
        try:
            self.current_file = filepath
            self.clear_pdf_properties()
        finally:
            self.current_file = old_current
            
    def clear_word_properties(self, filepath=None):
        try:
            target_file = filepath or self.current_file
            if target_file.lower().endswith('.docx'):
                doc = docx.Document(target_file)
                
                # 清除核心属性
                props = doc.core_properties
                props.title = None
                props.subject = None
                props.author = None
                props.category = None
                props.keywords = None
                props.comments = None
                props.last_modified_by = None
                
                # 保存到新文件
                temp_path = target_file + '.tmp'
                doc.save(temp_path)
                
                # 替换原文件
                os.replace(temp_path, target_file)
            else:
                # 对于DOC文件，在macOS/Linux上不提供清除功能
                raise Exception("DOC格式在macOS/Linux上不支持属性清除")
                
        except Exception as e:
            raise Exception(f"清除Word属性失败: {str(e)}")
            
    def clear_word_properties_file(self, filepath):
        """清除指定Word文件的属性"""
        old_current = self.current_file
        try:
            self.current_file = filepath
            self.clear_word_properties()
        finally:
            self.current_file = old_current
            
    def format_file_size(self, size):
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.2f} {unit}"
            size /= 1024
        return f"{size:.2f} TB"
        
    def format_timestamp(self, timestamp):
        import datetime
        return datetime.datetime.fromtimestamp(timestamp).strftime('%Y-%m-%d %H:%M:%S')
        
    def get_file_attributes(self, attrs):
        if IS_WINDOWS:
            attributes = []
            if attrs & 0x1: attributes.append("只读")
            if attrs & 0x2: attributes.append("隐藏")
            if attrs & 0x4: attributes.append("系统")
            if attrs & 0x10: attributes.append("目录")
            if attrs & 0x20: attributes.append("存档")
            return ", ".join(attributes) if attributes else "无特殊属性"
        else:
            return "平台不支持文件属性检查"
            
    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = FilePropertiesManager()
    app.run()